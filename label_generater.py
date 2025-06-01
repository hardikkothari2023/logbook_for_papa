import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

FILE_NAME = "contacts.xlsx"
OUTPUT_DOCX = "Labels_print.docx"

def load_data():
    if not os.path.exists(FILE_NAME):
        print("❌ No contact data found.")
        return pd.DataFrame(columns=["Name", "Phone", "Address", "City"])
    return pd.read_excel(FILE_NAME)

def get_contacts(df):
    print("\n--- 🖨️ Choose Contacts to Print ---")
    print("1. Print All Contacts")
    print("2. Filter by Name")
    print("3. Filter by City")
    choice = input("Choose option (1-3): ")

    if choice == "1":
        return df
    elif choice == "2":
        name = input("Enter name or part of it: ").strip().lower()
        return df[df["Name"].str.lower().str.contains(name)]
    elif choice == "3":
        city = input("Enter city: ").strip().lower()
        return df[df["City"].str.lower().str.contains(city)]
    else:
        print("❌ Invalid choice. Showing all.")
        return df

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # thicker border
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        borders.append(border)
    tcPr.append(borders)

def create_word_labels(data):
    if data.empty:
        print("📭 No contacts to print.")
        return

    doc = Document()

    # Margins for tighter layout
    section = doc.sections[0]
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    col_width = Inches(2.4)

    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    row_cells = None
    for i, (_, row) in enumerate(data.iterrows()):
        if i % 3 == 0:
            row_cells = table.add_row().cells

        label = (
            f"💌 {row['Name']}\n"
            f"📞 {row['Phone']}\n"
            f"🏠 {row['Address']}\n"
            f"🌆 {row['City']}"
        )

        para = row_cells[i % 3].paragraphs[0]
        run = para.add_run(label)
        run.font.name = 'Georgia'
        run.font.size = Pt(10.5)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.line_spacing = 1.1
        set_cell_border(row_cells[i % 3])

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Stylish labels saved as: {OUTPUT_DOCX}")

def main():
    df = load_data()
    if df.empty:
        return
    selected = get_contacts(df)
    create_word_labels(selected)

if __name__ == "__main__":
    main()
